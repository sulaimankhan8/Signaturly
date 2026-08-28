import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_user_guide():
    doc = docx.Document()

    # Set page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette: Signaturly Pro Brand (Dark Navy #090A0F, Red #EF4444, Gold #FACC15, Slate #1F2937, Light #F3F4F6)
    COLOR_PRIMARY = RGBColor(239, 68, 68)     # Red #EF4444
    COLOR_SECONDARY = RGBColor(250, 204, 21)  # Gold #FACC15
    COLOR_DARK = RGBColor(17, 24, 39)         # Dark Navy #111827
    COLOR_MUTED = RGBColor(107, 114, 128)     # Muted Gray #6B7280

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # =========================================================================
    # TITLE & HEADER SECTION
    # =========================================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("SIGNATURLY PRO")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(32)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("COMPREHENSIVE END-TO-END USER GUIDE & PAGE-BY-PAGE MANUAL")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_DARK

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(20)
    run_div = p_div.add_run("──────────────────────────────────────────────────────────────────────────")
    run_div.font.color.rgb = COLOR_MUTED
    run_div.font.size = Pt(9)

    # =========================================================================
    # SECTION 1: EXECUTIVE SUMMARY & LEGAL COMPLIANCE
    # =========================================================================
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Executive Summary & Legal Compliance Framework")
    run_h1.font.color.rgb = COLOR_PRIMARY
    run_h1.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Signaturly Pro is an enterprise electronic signature and document workflow platform built for maximum legal enforceability, security, and intuitive user experience. The application enforces complete compliance with major global e-signature acts:")

    # Compliance Table
    table_comp = doc.add_table(rows=4, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_comp.rows[0].cells
    hdr_cells[0].text = "Statutory Regulation"
    hdr_cells[1].text = "Jurisdiction"
    hdr_cells[2].text = "Compliance & Verification Mechanism"
    
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "111827")
        p_hdr = cell.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p_hdr.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

    rows_data = [
        ("US ESIGN Act (15 U.S.C. § 7001)", "United States (Federal)", "Intent to sign recording, electronic delivery consent gate, SHA-256 tamper seal, immutable audit certificate."),
        ("Section 10A Indian IT Act 2000", "India (National)", "Legal validity of electronic agreements, Email OTP recipient verification, timestamped IP audit trail."),
        ("EU eIDAS Regulation (No 910/2014)", "European Union", "Advanced Electronic Signature (AES) standards, cryptographic checksums, biometric signature canvas logging.")
    ]

    for r_idx, (col0, col1, col2) in enumerate(rows_data, start=1):
        row_cells = table_comp.rows[r_idx].cells
        row_cells[0].text = col0
        row_cells[1].text = col1
        row_cells[2].text = col2
        bg = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # =========================================================================
    # SECTION 2: PAGE-BY-PAGE DETAILED APPLICATION BREAKDOWN
    # =========================================================================
    h1_2 = doc.add_heading(level=1)
    run_h1_2 = h1_2.add_run("2. Comprehensive Page-by-Page Application Breakdown")
    run_h1_2.font.color.rgb = COLOR_PRIMARY
    run_h1_2.font.bold = True

    pages_data = [
        {
            "route": "/ & /landing",
            "name": "Home Redirect & Marketing Landing Page",
            "purpose": "Primary marketing hub showcasing Signaturly Pro capabilities, legal compliance trust badges, interactive demo previews, features, and quick navigation.",
            "features": [
                "Hero section with CTA buttons for instant account creation and user guide.",
                "Feature highlight cards detailing legal audit trails, prebuilt contracts, and OTP authorization.",
                "Interactive document security showcase with SHA-256 hash verification previews."
            ]
        },
        {
            "route": "/login",
            "name": "User Sign In Vault",
            "purpose": "Secure authentication portal for existing users to access their document vault and contract workspace.",
            "features": [
                "Email and Password input fields with automatic email normalization (.trim().toLowerCase()).",
                "Password visibility toggle and 'Forgot Password?' recovery link.",
                "JWT Access & HTTP-only Refresh Token authentication flow."
            ]
        },
        {
            "route": "/register",
            "name": "Account Registration & Mandatory Consent Gate",
            "purpose": "New user onboarding portal with built-in mandatory Terms & Conditions legal consent modal.",
            "features": [
                "Full Name, Email Address, and Password registration fields.",
                "Non-bypassable TermsConsentModal displaying ESIGN, IT Act 2000, and eIDAS disclosures.",
                "Dual statutory checkboxes requiring explicit assent before enabling the 'Accept & Proceed' action button."
            ]
        },
        {
            "route": "/forgot-password & /reset-password",
            "name": "Password Recovery & Reset Portal",
            "purpose": "Self-service password recovery workflow utilizing 1-hour expiration UUID tokens.",
            "features": [
                "Email address submission endpoint triggering automated password reset links.",
                "Secure token validation and password complexity check (min. 6 characters)."
            ]
        },
        {
            "route": "/dashboard",
            "name": "Document Vault & Workspace Dashboard",
            "purpose": "Central management console for executed agreements, pending signers, drafts, and document metrics.",
            "features": [
                "Vault KPI metrics cards (Total Documents, In Progress, Completed, Drafts).",
                "Filter tabs (All Docs, In Progress, Completed, Drafts, Declined/Void).",
                "Search bar filtering documents by title in real-time.",
                "Quick action buttons: 'Upload PDF', 'Use Template', 'Audit Certificate Download'."
            ]
        },
        {
            "route": "/upload",
            "name": "PDF Document Upload & Processing Engine",
            "purpose": "Portal for uploading custom PDF contracts into the Signaturly processing pipeline.",
            "features": [
                "Drag-and-drop PDF file upload area supporting files up to 50MB.",
                "Automatic page count parsing, PDF thumbnail preview, and title extraction."
            ]
        },
        {
            "route": "/assign/:pdfId",
            "name": "Visual Canvas Field Assignment Studio",
            "purpose": "Interactive studio for placing signature fields, dates, text inputs, and checkboxes onto PDF pages.",
            "features": [
                "Drag-and-drop visual field elements: Signature, Initials, Date, Text, Checkbox.",
                "Multi-recipient role assignment with customizable color-coding.",
                "Dynamic DPI coordinate mapper converting browser canvas pixels to standard 72 DPI PDF coordinates."
            ]
        },
        {
            "route": "/send/:pdfId",
            "name": "Recipient Dispatch & Security Settings",
            "purpose": "Configuration screen for defining signers, signing order, email OTP, and expiration dates.",
            "features": [
                "Multi-signer configuration with sequential or parallel signing order enforcement.",
                "Email OTP security toggle requiring 6-digit passcode verification before signing.",
                "Document expiration calendar picker and automated reminder schedule toggle."
            ]
        },
        {
            "route": "/sign/:token",
            "name": "Interactive Recipient Signing Portal",
            "purpose": "Public or authenticated portal for signers to execute documents via unique secure tokens.",
            "features": [
                "Email OTP verification step for protected documents.",
                "Interactive signature studio supporting drawn signatures, typed signatures (custom fonts), and uploaded signature images.",
                "Field completion tracking bar enforcing mandatory fields before final submission.",
                "Post-signing SHA-256 PDF flattening and Audit Certificate generation."
            ]
        },
        {
            "route": "/editor/:pdfId",
            "name": "PDF Document Viewer & Audit Inspector",
            "purpose": "Read-only inspection suite for viewing executed documents and audit trail histories.",
            "features": [
                "Multi-page high-resolution PDF canvas viewer.",
                "Audit log timeline detailing recipient IP addresses, timestamps, and hash changes."
            ]
        },
        {
            "route": "/templates",
            "name": "Pre-built Statutory Contract Template Suite",
            "purpose": "Library of 14 pre-built, legally compliant contract templates ready for instant dispatch.",
            "features": [
                "14 Pre-built contracts: NDA, Offer Letter, Contractor, Lease, Sales SOW, MSA, IP Assignment, SaaS License, Commercial Lease, Partnership, Bill of Sale, Promissory Note, Settlement Release, Board Resolution.",
                "Category filters (Legal & Corporate, Real Estate, HR, Sales, Consulting, Financial).",
                "Instant PDF preview and single-click 'Use Template' action."
            ]
        },
        {
            "route": "/templates/bulk",
            "name": "CSV Bulk Dispatch Studio",
            "purpose": "Batch campaign dispatch generator for sending contracts to hundreds of recipients simultaneously via CSV upload.",
            "features": [
                "CSV template downloader and batch data parser.",
                "Field mapping column selector matching CSV headers to contract recipient roles.",
                "Automated campaign execution with live progress tracker."
            ]
        },
        {
            "route": "/templates/edit/:templateId & /use/:templateId",
            "name": "Template Customization & Instantiation Studio",
            "purpose": "Interface for customizing template field positions and instantiating templates for signing.",
            "features": [
                "Visual field editor allowing pre-placement of signature anchors.",
                "Recipient field pre-filling and dynamic contract generation."
            ]
        },
        {
            "route": "/signature-remover",
            "name": "Signature Studio & Background Eraser Tool",
            "purpose": "Utility tool for removing background noise and making physical handwritten signature uploads transparent.",
            "features": [
                "Threshold-based background transparency eraser.",
                "Signature crop, rotation, and high-contrast enhancement."
            ]
        },
        {
            "route": "/settings",
            "name": "User Profile & Security Preferences",
            "purpose": "User management screen for profile settings, password changes, and API preferences.",
            "features": [
                "Profile name and email updates.",
                "Password modification form.",
                "Terms acceptance status indicator displaying UTC timestamp and recorded IP address."
            ]
        },
        {
            "route": "/verify",
            "name": "Public Cryptographic Audit Verification Portal",
            "purpose": "Publicly accessible verification tool for validating document authenticity and tamper integrity.",
            "features": [
                "PDF file drag-and-drop or SHA-256 hash search bar.",
                "Independent verification engine checking pre/post signing checksums against database audit logs."
            ]
        },
        {
            "route": "/admin/login & /admin/dashboard",
            "name": "Superadmin Oversight & System Analytics",
            "purpose": "Admin console for monitoring platform performance, user accounts, and system health.",
            "features": [
                "Superadmin secret key authentication.",
                "Platform metrics (Total Users, Total PDFs Executed, Storage Utilization).",
                "User account management and document revocation controls."
            ]
        },
        {
            "route": "/userguide",
            "name": "Interactive Online User Documentation Suite",
            "purpose": "Built-in interactive documentation suite accessible directly within the application web interface.",
            "features": [
                "Searchable navigation sidebar with interactive feature walkthroughs.",
                "Code snippets, API references, and video guide embeds."
            ]
        },
        {
            "route": "/test",
            "name": "Font & Canvas Render Sandbox",
            "purpose": "Isolated testing lab for verifying custom signature font rendering and canvas flattening.",
            "features": [
                "Live font previewer testing signature font styles (Dancing Script, Great Vibes, Pacifico, Alex Brush).",
                "PDF-lib canvas rendering test suite."
            ]
        }
    ]

    for p_info in pages_data:
        h2 = doc.add_heading(level=2)
        run_h2 = h2.add_run(f"Route: {p_info['route']} ── {p_info['name']}")
        run_h2.font.color.rgb = COLOR_DARK
        run_h2.font.bold = True

        p_purp = doc.add_paragraph()
        run_p_lbl = p_purp.add_run("Purpose: ")
        run_p_lbl.bold = True
        p_purp.add_run(p_info["purpose"])

        p_feat_hdr = doc.add_paragraph()
        run_f_lbl = p_feat_hdr.add_run("Key Features & Functionality:")
        run_f_lbl.bold = True

        for feat in p_info["features"]:
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.add_run(feat)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Output path
    output_path = os.path.join("docs", "USER_GUIDE_MANUAL.docx")
    doc.save(output_path)
    print(f"SUCCESS: USER_GUIDE_MANUAL.docx generated successfully at: {output_path}")

if __name__ == "__main__":
    create_user_guide()
