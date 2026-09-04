import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_header_banner(title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Dark Blue
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(13)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(100, 116, 139) # Slate
    doc.add_paragraph()

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110) # Teal
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def add_p(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code(code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9") # Slate light
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()

# Document Content Generation
add_header_banner(
    "Signaturly Pro — Cloud Deployment, Docker & GCP Masterclass",
    "Complete Reference Guide & Executive Refresher for Enterprise Cloud Architecture"
)

add_h1("1. Executive Summary & Zero-Cost Architecture ($0.00 / Month)")
add_p("Signaturly Pro is architected to leverage the Always-Free Tiers of Google Cloud Platform (GCP), MongoDB Atlas, and Vercel Global Edge Network. The platform achieves enterprise-grade legal compliance (US ESIGN, EU eIDAS, IT Act 2000 Section 10A) with zero ongoing hosting expense.")

# Cost Table
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ["Component", "Provider & Tier", "Specs & Quotas", "Monthly Cost"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_shading(hdr_cells[i], "1E3A8A")
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data = [
    ("Backend API", "GCP Cloud Run (Serverless)", "2,000,000 requests/mo + Scale to 0", "$0.00"),
    ("Document Vault", "Google Cloud Storage (GCS)", "5 GB Standard Storage (us-central1)", "$0.00"),
    ("Container Registry", "GCP Artifact Registry", "500 MB Container Image Store", "$0.00"),
    ("Customer Web App", "Vercel Hobby Tier", "100 GB Bandwidth, Global Edge CDN", "$0.00"),
    ("Admin Portal", "Vercel Hobby Tier", "Isolated Governance UI on Edge", "$0.00"),
    ("Database", "MongoDB Atlas (M0)", "512 MB TLS Encrypted Cloud Cluster", "$0.00"),
    ("SSL Certificates", "Google & Vercel Managed", "Auto-renewing 2048-bit RSA / ECC", "$0.00"),
]

for row in data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val
        row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()

add_h1("2. Docker & Container Architecture Crash Course")
add_h2("The 4 Core Docker Concepts")
add_p("• Dockerfile: A plain text file containing step-by-step instructions for packaging your code, Node.js runtime, OS libraries, and dependencies (The Cooking Recipe).")
add_p("• Docker Image: An immutable, standalone package created from building the Dockerfile (The Baked Meal Template).")
add_p("• Artifact Registry: Google Cloud's cloud repository where your built images are stored (The Supermarket Shelf).")
add_p("• Docker Container: An active, running process created from an image in memory (Eating the Meal Now).")

add_h2("Multi-Stage Builds & Security Hardening")
add_p("The production Dockerfile in server/Dockerfile utilizes a 3-stage multi-stage architecture:")
add_p("1. Stage 1 (Base): Installs OpenSSL, curl, and Debian CA certificates needed for cryptographic SHA-256 PDF signatures and audit trails.")
add_p("2. Stage 2 (Dependencies): Runs npm ci --only=production to install exact dependency trees with zero development bloat.")
add_p("3. Stage 3 (Runner): Copies only production node_modules and code. Runs as USER node (non-root) to eliminate container breakout risks.")

add_h1("3. GCP Architecture & Components Breakdown")
add_p("• Google Cloud Run: Serverless Container execution platform. Scales to 0 instances when idle ($0.00 compute). Spawns in ~1.5s on cold starts and dynamically assigns PORT 8080.")
add_p("• Google Cloud Storage (GCS): Persistent cloud object storage. Since Cloud Run containers are ephemeral (disk clears on shutdown), all PDF agreements and signed certificates are vaulted in GCS (5GB Always-Free Tier). Time-limited Signed URLs (15 min) allow secure document streaming without exposing the bucket.")
add_p("• Artifact Registry: Google's standardized OCI container registry storing Docker images in us-central1.")
add_p("• IAM & Service Accounts: Robot identities with granular roles (roles/run.admin, roles/artifactregistry.writer, roles/storage.admin).")

add_h1("4. Step-by-Step GCP Deployment Playbook (Commands & Flags)")

add_h2("Command 1: Enable GCP Cloud APIs")
add_code("""gcloud services enable run.googleapis.com artifactregistry.googleapis.com cloudbuild.googleapis.com storage.googleapis.com iamcredentials.googleapis.com logging.googleapis.com --project=project-9aac115a-73f9-4e8a-9a7""")
add_p("Enables Cloud Run serverless engine, Artifact Registry container store, Cloud Build compiler, GCS vault, Workload Identity OIDC, and Cloud Logging.")

add_h2("Command 2: Grant Cloud Build Service Account IAM Roles")
add_code("""gcloud projects add-iam-policy-binding project-9aac115a-73f9-4e8a-9a7 --member="serviceAccount:774331940137-compute@developer.gserviceaccount.com" --role="roles/storage.admin"
gcloud projects add-iam-policy-binding project-9aac115a-73f9-4e8a-9a7 --member="serviceAccount:774331940137-compute@developer.gserviceaccount.com" --role="roles/artifactregistry.writer"
gcloud projects add-iam-policy-binding project-9aac115a-73f9-4e8a-9a7 --member="serviceAccount:774331940137-compute@developer.gserviceaccount.com" --role="roles/logging.logWriter\"""")
add_p("Grants the default Google Cloud Build service account permission to upload build archives, write images to Artifact Registry, and stream build logs.")

add_h2("Command 3: Create Always-Free GCS Storage Bucket")
add_code("""gcloud storage buckets create gs://signaturly-vault-project-9aac115a-73f9-4e8a-9a7 --location=us-central1 --default-storage-class=STANDARD --uniform-bucket-level-access""")
add_p("Creates private storage bucket in us-central1 (qualifying for 5GB free tier forever) with uniform bucket-level IAM enforcement.")

add_h2("Command 4: Create Artifact Registry Docker Repository")
add_code("""gcloud artifacts repositories create signaturly-repo --repository-format=docker --location=us-central1 --description="Signaturly Docker Repository" --project=project-9aac115a-73f9-4e8a-9a7""")
add_p("Creates the named OCI Docker repository colocated in us-central1 for zero network egress latency.")

add_h2("Command 5: Build Image with Cloud Build")
add_code("""cd ~/Signaturly/server
gcloud builds submit --tag us-central1-docker.pkg.dev/project-9aac115a-73f9-4e8a-9a7/signaturly-repo/signaturly-api:latest""")
add_p("Compresses local code, uploads archive to GCP, compiles Docker container on Google infrastructure, and registers the image.")

add_h2("Command 6: Deploy to Google Cloud Run")
add_code("""gcloud run deploy signaturly-api \\
  --image us-central1-docker.pkg.dev/project-9aac115a-73f9-4e8a-9a7/signaturly-repo/signaturly-api:latest \\
  --platform managed \\
  --region us-central1 \\
  --allow-unauthenticated \\
  --min-instances 0 \\
  --max-instances 1 \\
  --memory 512Mi \\
  --set-env-vars="NODE_ENV=production,CORS_ORIGIN=*,MONGO_URI=mongodb+srv://suleman111111111111111_db_user:WOCfS74qN0fbJcc8@cluster0.0e0cvkg.mongodb.net/?appName=Cluster0,JWT_ACCESS_SECRET=access_secret_123,JWT_REFRESH_SECRET=refresh_secret_456,ADMIN_SECRET_KEY=signaturly-superadmin-secret,STORAGE_PROVIDER=gcs,GCP_PROJECT_ID=project-9aac115a-73f9-4e8a-9a7,GCS_BUCKET_NAME=signaturly-vault-project-9aac115a-73f9-4e8a-9a7,GCS_SIGNED_URL_EXPIRES=900,SMTP_HOST=smtp.gmail.com,SMTP_PORT=465,SMTP_USER=suleman111111111111111@gmail.com,SMTP_PASS=vgsatqmnxjowucfs,SMTP_FROM=suleman111111111111111@gmail.com" \\
  --project project-9aac115a-73f9-4e8a-9a7""")
add_p("Deploys container to serverless Cloud Run with min-instances 0 (scale-to-zero) and max-instances 1 (hard billing ceiling).")

add_h1("5. Keyless GitHub Actions CI/CD (Workload Identity Federation)")
add_p("Instead of using static, insecure JSON private keys, GitHub Actions connects to Google Cloud using OpenID Connect (OIDC) via Workload Identity Federation (WIF).")

add_h2("WIF Setup Commands:")
add_code("""# 1. Enable IAM Credentials API
gcloud services enable iamcredentials.googleapis.com --project=project-9aac115a-73f9-4e8a-9a7

# 2. Create Workload Identity Pool
gcloud iam workload-identity-pools create "github-pool" --project="project-9aac115a-73f9-4e8a-9a7" --location="global" --display-name="GitHub Pool"

# 3. Create OIDC Provider
gcloud iam workload-identity-pools providers create-oidc "github-provider" \\
  --project="project-9aac115a-73f9-4e8a-9a7" \\
  --location="global" \\
  --workload-identity-pool="github-pool" \\
  --display-name="GitHub Provider" \\
  --attribute-mapping="google.subject=assertion.sub,attribute.repository=assertion.repository" \\
  --attribute-condition="assertion.repository == 'sulaimankhan8/Signaturly'" \\
  --issuer-uri="https://token.actions.githubusercontent.com"

# 4. Bind GitHub Repo to Service Account
gcloud iam service-accounts add-iam-policy-binding "github-actions-deployer@project-9aac115a-73f9-4e8a-9a7.iam.gserviceaccount.com" \\
  --project="project-9aac115a-73f9-4e8a-9a7" \\
  --role="roles/iam.workloadIdentityUser" \\
  --member="principalSet://iam.googleapis.com/projects/774331940137/locations/global/workloadIdentityPools/github-pool/attribute.repository/sulaimankhan8/Signaturly\"""")

add_h2("Sequential GitHub Actions Workflow (.github/workflows/deploy-backend.yml):")
add_p("1. Push Trigger: Triggers on git push origin main when files in server/** change.")
add_p("2. Ubuntu Runner: GitHub allocates a clean virtual environment.")
add_p("3. Checkout Repository: Fetches source code.")
add_p("4. Google Auth Action: Authenticates via WIF and receives a short-lived token (1hr).")
add_p("5. Cloud SDK Setup: Configures gcloud on the runner.")
add_p("6. Docker Authentication: Runs gcloud auth configure-docker us-central1-docker.pkg.dev.")
add_p("7. Build & Push: Builds Docker image tagged with commit SHA and pushes to Artifact Registry.")
add_p("8. Deploy to Cloud Run: Updates Cloud Run service to newly pushed image with zero downtime.")

add_h1("6. Frontend Deployment on Vercel & SPA Rewrites (vercel.json)")
add_p("• React Single Page Applications (SPAs) are compiled static bundles. Vercel serves them globally over edge CDN nodes.")
add_p("• SPA Problem & vercel.json: Without rewrites, refreshing browser routes like /dashboard or /verify results in 404 errors. vercel.json routes all requests to index.html, letting React Router handle routing in-browser.")
add_p("• Setup: Link repository in Vercel > set Root Directory (client for user app, admin for governance portal) > set VITE_API_BASE_URL to https://signaturly-api-774331940137.us-central1.run.app.")

add_h1("7. Superadmin Governance Portal Credentials")
add_p("• Admin Portal: client/admin on Vercel or localhost:5174")
add_p("• Admin Email: admin@signaturly.com")
add_p("• Admin Password: Admin@123456")
add_p("• Admin Secret Key: signaturly-superadmin-secret")

add_h1("8. Disaster Recovery & Clean Migration Playbook")
add_p("To delete all cloud resources:")
add_code("""gcloud run services delete signaturly-api --region=us-central1 --quiet
gcloud artifacts repositories delete signaturly-repo --location=us-central1 --quiet
gcloud storage rm --recursive gs://signaturly-vault-project-9aac115a-73f9-4e8a-9a7""")
add_p("To deploy on a brand-new GCP account/project:")
add_p("1. Set active project: gcloud config set project NEW_PROJECT_ID")
add_p("2. Run Commands 1 through 6 in Section 4.")
add_p("3. Update VITE_API_BASE_URL in Vercel to the new Cloud Run URL.")

# Save Document
out_path = os.path.abspath("docs/CLOUD_DEPLOYMENT_DOCKER_GCP_MASTERCLASS.docx")
doc.save(out_path)
print("SUCCESS: Generated docx at:", out_path)
